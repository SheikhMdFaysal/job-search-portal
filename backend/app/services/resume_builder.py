from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.shared import Pt

from app.core.config import settings
from app.integrations.openai_client import chat_json
from app.models import Application, Job, Profile


def _safe_slug(value: str) -> str:
    return re.sub(r"[^a-zA-Z0-9]+", "-", value).strip("-").lower()[:80]


def _fallback_resume(profile: Profile, job: Job) -> dict[str, Any]:
    resume = profile.master_resume
    skills = resume.get("skills", [])
    if isinstance(skills, str):
        skills = [part.strip() for part in skills.split(",")]
    return {
        "name": resume.get("name", "Candidate"),
        "headline": f"Candidate for {job.title}",
        "summary": [
            f"Analytics-focused professional targeting {job.title} roles with experience in data, reporting, and stakeholder support."
        ],
        "skills": skills[:18],
        "experience": resume.get("experience", []),
        "education": resume.get("education", []),
        "certifications": resume.get("certifications", []),
    }


async def tailor_resume(profile: Profile, application: Application) -> dict[str, Any]:
    job = application.job
    fallback = _fallback_resume(profile, job)
    return await chat_json(
        {
            "master_resume": profile.master_resume,
            "job_description": job.description,
            "job_title": job.title,
            "company": job.company,
            "rules": [
                "ATS friendly only",
                "No tables, columns, graphics, icons, or text boxes",
                "Use standard sections: Summary, Experience, Education, Skills, Certifications",
                "Prefer quantified bullets when supported by the master resume",
                "Do not invent employers, degrees, dates, or certifications",
            ],
            "return_schema": fallback,
        },
        fallback=fallback,
    )


def render_resume_docx(application: Application, tailored: dict[str, Any]) -> str:
    storage = Path(settings.document_storage_dir)
    storage.mkdir(parents=True, exist_ok=True)
    job = application.job
    filename = f"{application.id}-{_safe_slug(job.company)}-{_safe_slug(job.title)}-resume.docx"
    path = storage / filename

    doc = DocxDocument()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    doc.add_heading(str(tailored.get("name", "Candidate")), level=0)
    headline = tailored.get("headline")
    if headline:
        doc.add_paragraph(str(headline))

    for section_name in ["Summary", "Skills", "Experience", "Education", "Certifications"]:
        key = section_name.lower()
        value = tailored.get(key)
        if not value:
            continue
        doc.add_heading(section_name, level=1)
        if isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    title = item.get("title") or item.get("role") or item.get("degree") or item.get("name")
                    org = item.get("company") or item.get("school") or item.get("issuer")
                    if title or org:
                        doc.add_paragraph(" - ".join(str(part) for part in [title, org] if part))
                    bullets = item.get("bullets") or item.get("details") or []
                    for bullet in bullets if isinstance(bullets, list) else [bullets]:
                        doc.add_paragraph(str(bullet), style="List Bullet")
                else:
                    doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph(str(value))

    doc.core_properties.comments = json.dumps({"application_id": application.id, "job_id": job.id})
    doc.save(path)
    return str(path)
